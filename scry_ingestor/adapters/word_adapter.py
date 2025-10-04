"""Unstructured adapter for Word documents (.docx format only)."""

import io
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pydantic import ValidationError as PydanticValidationError

from ..exceptions import CollectionError, ConfigurationError
from ..schemas.payload import ValidationResult
from ..schemas.transformations import WordTransformationConfig
from ..utils.file_readers import read_binary_file, resolve_binary_read_options
from .base import BaseAdapter


class WordAdapter(BaseAdapter):
    """
    Adapter for collecting and processing unstructured text from Word documents.

    Supports:
    - .docx files (Office 2007+ / Office Open XML format)
    - Text extraction from paragraphs
    - Metadata extraction (author, title, subject, keywords, timestamps)
    - Table extraction (optional)

    Note: This adapter only supports .docx files. Legacy .doc files (Office 97-2003)
    are not supported. Users with .doc files should convert them to .docx first using:
    - Microsoft Word: File > Save As > .docx
    - LibreOffice: File > Save As > .docx
    - Online converters (e.g., CloudConvert, Zamzar)
    """

    SUPPORTED_FORMAT = ".docx"

    def __init__(self, config: dict[str, Any]):
        super().__init__(config)
        try:
            self._transformation = WordTransformationConfig.model_validate(
                config.get("transformation") or {}
            )
        except PydanticValidationError as exc:
            raise ConfigurationError(
                f"Invalid Word transformation configuration: {exc}"
            ) from exc

    async def collect(self) -> Any:
        """
        Collect raw data from Word document (.docx only).

        Returns:
            DocxDocument object containing document structure and content

        Raises:
            CollectionError: If document collection fails or unsupported format
        """
        source_type = self.config.get("source_type", "file")

        try:
            if source_type == "file":
                file_path = self.config.get("path")
                if not file_path:
                    raise CollectionError("File path not provided in config")

                path = Path(file_path)
                if not path.exists():
                    raise CollectionError(f"File not found: {file_path}")

                file_format = path.suffix.lower()
                if file_format != self.SUPPORTED_FORMAT:
                    raise CollectionError(
                        f"Unsupported file type: {file_format}. "
                        f"Only .docx files are supported. "
                        f"For .doc files, please convert to .docx format first."
                    )

                chunk_size, max_bytes = resolve_binary_read_options(
                    self.config.get("read_options")
                )
                data_bytes = await self._run_in_thread(
                    read_binary_file,
                    file_path,
                    chunk_size=chunk_size,
                    max_bytes=max_bytes,
                )
                buffer = io.BytesIO(data_bytes)
                return await self._run_in_thread(DocxDocument, buffer)

            else:
                raise CollectionError(f"Unsupported source type: {source_type}")

        except CollectionError:
            raise
        except OSError as e:
            raise CollectionError(f"Failed to read Word document: {e}")
        except Exception as e:
            raise CollectionError(f"Failed to parse Word document: {e}") from e

    async def validate(self, raw_data: Any) -> ValidationResult:
        """
        Validate the Word document structure and content.

        Args:
            raw_data: DocxDocument object from collect()

        Returns:
            ValidationResult with quality metrics
        """
        errors = []
        warnings = []
        metrics = {}

        try:
            doc = raw_data

            # Count paragraphs
            paragraph_count = len(doc.paragraphs)
            metrics["paragraph_count"] = paragraph_count

            # Count tables
            table_count = len(doc.tables)
            metrics["table_count"] = table_count

            # Extract text and count
            text_content = "\n".join([p.text for p in doc.paragraphs])
            text_length = len(text_content.strip())
            metrics["text_length_chars"] = text_length
            metrics["word_count"] = len(text_content.split())

            # Validation rules from config
            min_paragraphs = self.config.get("validation", {}).get("min_paragraphs", 0)
            min_words = self.config.get("validation", {}).get("min_words", 0)
            allow_empty = self.config.get("validation", {}).get("allow_empty", False)

            # Check for empty document
            if text_length == 0 and not allow_empty:
                errors.append("Document contains no text content")

            # Check minimum paragraphs
            if paragraph_count < min_paragraphs:
                errors.append(
                    f"Document has {paragraph_count} paragraphs, "
                    f"minimum required: {min_paragraphs}"
                )

            # Check minimum words
            word_count = metrics["word_count"]
            if word_count < min_words:
                errors.append(
                    f"Document has {word_count} words, minimum required: {min_words}"
                )

            # Warnings for potentially problematic documents
            if paragraph_count == 0 and table_count > 0:
                warnings.append("Document contains only tables, no text paragraphs")

            is_valid = len(errors) == 0

        except Exception as e:
            errors.append(f"Validation error: {str(e)}")
            is_valid = False

        return ValidationResult(
            is_valid=is_valid,
            errors=errors,
            warnings=warnings,
            metrics=metrics,
        )

    async def transform(self, raw_data: Any) -> dict[str, Any]:
        """
        Transform Word document into standardized text format.

        Args:
            raw_data: DocxDocument object from collect()

        Returns:
            Dictionary with extracted text, metadata, and optional table data
        """
        transformation_config = self._transformation
        doc = raw_data

        # Extract text from paragraphs
        paragraphs_raw = [p.text for p in doc.paragraphs]
        if transformation_config.strip_whitespace:
            paragraphs = [text.strip() for text in paragraphs_raw if text and text.strip()]
        else:
            paragraphs = [text for text in paragraphs_raw if text]

        # Join paragraphs
        separator = transformation_config.paragraph_separator
        full_text = separator.join(paragraphs)

        # Extract core properties (metadata)
        metadata = {}
        if transformation_config.extract_metadata:
            core_props = doc.core_properties
            metadata = {
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": (
                    core_props.modified.isoformat() if core_props.modified else None
                ),
            }

        # Extract tables
        tables_data = []
        if transformation_config.extract_tables:
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    table_rows.append([cell.text for cell in row.cells])
                tables_data.append(table_rows)

        result = {
            "text": full_text,
            "paragraph_count": len(paragraphs),
            "metadata": metadata,
        }

        if tables_data:
            result["tables"] = tables_data

        return result
